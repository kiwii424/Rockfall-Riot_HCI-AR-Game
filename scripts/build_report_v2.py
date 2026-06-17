# -*- coding: utf-8 -*-
"""Build the restructured IEEE-style report (v2) as .docx.

Applies the research-paper-writing skill workflow:
  - one message per paragraph, message in first sentence
  - Method written as per-pillar (motivation -> design -> advantage)
  - Experiments as setup -> validation -> profiling -> honest limits
  - claim-evidence map + five-dimension self-review appended

All quantitative numbers are the verified ones from the codebase/eval/benchmark.
Two factual corrections vs the original draft are baked in:
  - single-hand input (max_num_hands=1), not "two hands"
  - author block has two authors, no placeholder
"""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

OUT = Path.home() / "Downloads" / "Rockfall_Riot_Report_v2.docx"

doc = Document()

# ---- base style ----
normal = doc.styles["Normal"]
normal.font.name = "Times New Roman"
normal.font.size = Pt(10)


def para(text, *, bold=False, italic=False, align=None, size=None, space_after=6, color=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    if size:
        r.font.size = Pt(size)
    if color:
        r.font.color.rgb = RGBColor(*color)
    if align:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    return p


def heading(text, level=1):
    sizes = {1: 12, 2: 11}
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(sizes.get(level, 11))
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    return p


def body(text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return p


# ================= TITLE / AUTHORS =================
para("Rockfall Riot: A Hand-Gesture-Controlled Augmented-Reality Rhythm "
     "Game with Music-Synchronized Spawning",
     bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=16, space_after=8)

para("Sung-Yu Hong (洪崧祐, 113511216)    "
     "Yu-Chi Huang (黃淯琪, 111511157)",
     align=WD_ALIGN_PARAGRAPH.CENTER, size=11, space_after=2)
para("Group 4 — Machine Learning Project",
     align=WD_ALIGN_PARAGRAPH.CENTER, size=10, space_after=2)
para("Department of Electrical and Computer Engineering, "
     "National Yang Ming Chiao Tung University, Hsinchu, Taiwan",
     italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=10, space_after=10)

# ================= ABSTRACT (Version 3: multi-contribution + advantage) =====
heading("Abstract", 1)
body(
    "We present Rockfall Riot, a real-time augmented-reality (AR) rhythm game in "
    "which a standard webcam replaces the game controller and the player's bare "
    "hand is the sole input device. The system integrates three machine-learning "
    "components into one low-latency gameplay loop, and each component is paired "
    "with an explicit engineering advantage. First, markerless perception uses "
    "Google's pretrained MediaPipe Hands pipeline, which returns 21 2.5-D "
    "landmarks per hand and lets us avoid any depth sensor. Second, gesture "
    "recognition uses a Random Forest trained on a self-collected dataset of 807 "
    "samples over a 42-dimensional wrist-relative, scale-normalized feature "
    "vector, which makes the classifier fast to train and invariant to hand "
    "position and distance. Third, music-synchronized spawning is driven by a "
    "pretrained Temporal Convolutional Network (TCN) beat tracker decoded by a "
    "dynamic Bayesian network (DBN), which yields beat timestamps and strengths "
    "that set obstacle density and fall speed. To keep a controller-free "
    "interface playable we add a multi-layer fallback architecture and a "
    "temporal-smoothing scheme that suppress jitter, together with dynamic "
    "difficulty adjustment. On a three-class dataset the gesture classifier "
    "reaches 98.14% (±1.47%) five-fold cross-validation accuracy and 99.38% "
    "hold-out accuracy; we stress that this figure is optimistic given the small, "
    "imbalanced data. The full perception pipeline runs at roughly 30 frames per "
    "second with about 33 ms end-to-end latency, and the beat tracker runs as the "
    "default backend with graceful degradation to lighter algorithms when its "
    "dependencies are absent."
)
para("Index Terms—Gesture recognition, beat tracking, augmented reality, "
     "human–computer interaction, random forest, temporal convolutional "
     "network, dynamic Bayesian network.", italic=True, size=9)

# ================= I. INTRODUCTION =================
heading("I.  Introduction", 1)
body(
    "Controller-free interaction driven by a single RGB camera is the setting "
    "this paper targets. Vision-based hand tracking now localizes the hand "
    "skeleton in real time on commodity hardware, and rhythm games demand a tight "
    "temporal coupling between player action and music. Rockfall Riot joins these "
    "two threads: the player faces a webcam, falling rocks are spawned in time "
    "with the music, and the player destroys or catches them using hand gestures "
    "alone."
)
body(
    "Building such a system raises three concrete challenges, and each one maps to "
    "a component of our pipeline. First, a real-time gesture controller must be "
    "robust, because per-frame predictions from a vision model are noisy and brief "
    "detection losses must not make the on-screen cursor disappear. Second, game "
    "events must lock onto the musical beat, which requires reliable beat and "
    "onset estimation from arbitrary audio. Third, the difficulty must adapt to "
    "the player so that the game is neither trivial nor frustrating."
)
body(
    "A webcam-only interface is attractive for accessibility and cost, because it "
    "removes specialized controllers and depth sensors and runs on hardware "
    "players already own. A Fruit-Ninja-style rhythm game is a useful testbed for "
    "this interface because it exercises the whole pipeline at once — accurate "
    "gesture classification, low latency, spatial precision for slicing, and tight "
    "audio synchronization — so a weakness in any single component is "
    "immediately visible to the player."
)
body("The contributions of this project are as follows.")
for c in [
    "An end-to-end, controller-free gesture pipeline that combines pretrained "
    "MediaPipe hand landmarks with a lightweight Random Forest classifier, backed "
    "by a multi-layer fallback (MediaPipe → OpenCV color tracking) that "
    "degrades gracefully when landmark detection fails.",
    "A temporal-smoothing scheme — a nine-frame majority vote with switching "
    "hysteresis and a short-horizon position hold — that converts noisy "
    "per-frame predictions into a stable control signal.",
    "A music-driven rhythm module backed by a runnable multi-backend beat "
    "pipeline (pretrained TCN+DBN → madmom → librosa → metronome) "
    "that maps beat timestamps and strengths to obstacle density and fall speed.",
    "A player-modeling layer that provides dynamic difficulty adjustment and a "
    "post-game skill profile rendered as a radar chart.",
]:
    p = doc.add_paragraph(c, style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

# ================= II. RELATED WORK =================
heading("II.  Related Work", 1)
heading("A.  Hand Tracking", 2)
body(
    "MediaPipe Hands [1], [2] is a two-stage pipeline that we use directly as our "
    "perception front end. A BlazePalm [3] detector first locates the palm and "
    "returns an oriented bounding box, and a landmark model then regresses 21 "
    "2.5-D keypoints inside the cropped region together with a hand-presence flag "
    "and a handedness label. Because the detector is re-run only when tracking is "
    "lost, the pipeline reaches real-time inference on commodity devices without "
    "any depth sensor."
)
heading("B.  Beat Tracking", 2)
body(
    "Beat tracking has shifted from signal-processing onset functions toward deep "
    "models, and our default tracker follows that line. Recurrent BLSTM trackers "
    "[5], [6] decoded by a dynamic Bayesian network (DBN) [7] long defined the "
    "state of the art. Davies and Böck [4] replaced the BLSTM with a Temporal "
    "Convolutional Network (TCN) [8] that performs dilated convolutions — the "
    "mechanism popularized by WaveNet [9] — achieving comparable accuracy with "
    "roughly one third of the weights and far faster training. We adopt a "
    "pretrained TCN–DBN tracker of this family as our primary beat estimator, "
    "with the classical dynamic-programming tracker of Ellis [14] as a lightweight "
    "fallback."
)
heading("C.  Gesture Classification on Landmarks", 2)
body(
    "Once a hand is reduced to a small vector of normalized landmark coordinates, "
    "the classification problem becomes low-dimensional and tabular, which favors "
    "ensemble methods. Random Forests [10] train quickly on small datasets, "
    "require no feature scaling, are robust to irrelevant features, and expose "
    "feature importances that reveal which landmarks drive a decision. A deep "
    "network would be data-hungry and harder to justify on a few hundred samples, "
    "so we favor a Random Forest for the gesture-recognition stage."
)
heading("D.  Dynamic Difficulty Adjustment", 2)
body(
    "Keeping a player in a state of flow is a long-standing goal of game design, "
    "and online difficulty adaptation is a common mechanism for it. We apply a "
    "lightweight version of this idea, driving the obstacle-spawn rate directly "
    "from a short-window hit/miss statistic rather than from a fixed level "
    "schedule."
)

# ================= III. METHOD (three-element per pillar) =================
heading("III.  System Architecture", 1)
body(
    "The system runs two concurrent real-time streams that feed a shared gameplay "
    "loop. A hand-gesture stream and a music/rhythm stream are merged so that "
    "input stability and the musical beat jointly determine which game events "
    "occur. Section III-A to III-D describe the gesture stream (perception, "
    "features, classification, smoothing), Section III-E the multi-layer "
    "robustness design, Section III-F the beat stream, and Section III-G to III-H "
    "the gameplay and player-modeling layers; the two streams and their merge "
    "point are summarized in Fig. 1."
)
heading("A.  Hand Tracking and Landmark Extraction", 2)
body(
    "Perception must turn a raw camera frame into a stable hand skeleton, which we "
    "delegate to MediaPipe Hands. Each frame is passed to the model, which returns "
    "the 21 landmarks of the detected hand as (x, y, relative depth) coordinates. "
    "We consume the 21 (x, y) image-plane coordinates and we track a single hand "
    "(num_hands = 1); the relative-depth channel is available but not used by the "
    "classifier."
)
heading("B.  Feature Engineering", 2)
body(
    "Raw image coordinates are not directly usable because they vary with hand "
    "position and camera distance, so we normalize each sample in two steps. The "
    "21 landmarks first yield 42 raw coordinate values. We then subtract the wrist "
    "landmark from every coordinate to obtain translation invariance, and divide "
    "all coordinates by the maximum absolute value in the sample to obtain scale "
    "invariance. The result is a 42-dimensional, wrist-relative, scale-normalized "
    "feature vector, which makes the classifier insensitive to where the hand "
    "appears and how far it is from the camera. For visualization we additionally "
    "project the features to two dimensions with PCA; the first two components "
    "capture 51.4% and 31.9% of the variance (83.3% jointly), and in this "
    "projection the open-palm “Stop” class forms a compact, "
    "well-separated cluster while “Sword” and “Fist” partially "
    "overlap — a structure that predicts the dominant error mode of "
    "Section IV."
)
heading("C.  Gesture Classification", 2)
body(
    "The classifier maps the normalized feature vector to one of three gestures: "
    "Sword (extended index finger), Fist (closed hand), and Stop (open palm). A "
    "Random Forest is well matched to the problem because the dataset is small, "
    "the features are tabular, training is fast, and the model exposes feature "
    "importances for inspection. Model selection uses five-fold cross-validation."
)
heading("D.  Real-Time Temporal Smoothing", 2)
body(
    "A frame-by-frame classifier flickers, because individual frames are "
    "occasionally misclassified and the cursor jumps, so we stabilize the output "
    "with three mechanisms. A nine-frame sliding window aggregates the most recent "
    "predictions so that a single bad frame cannot flip the decision. A 60% "
    "majority hysteresis requires a new label to occupy more than 60% of the "
    "window before the active gesture switches, which prevents oscillation when a "
    "gesture sits on a decision boundary. Finally, a position hold keeps the last "
    "valid hand position for up to eight frames when the hand is briefly "
    "undetected, so the cursor does not vanish during short dropouts."
)
heading("E.  Robustness via Multi-Layer Fallback", 2)
body(
    "Both perception stages degrade gracefully rather than failing outright, which "
    "keeps the controller-free interface usable across machines and conditions. "
    "For hand tracking, MediaPipe is the primary backend, and when landmarks are "
    "unavailable an OpenCV color-based blob tracker supplies a coarse but stable "
    "palm position; this fallback has the lowest precision but the highest "
    "stability, and its role is to prevent visible flicker, not to drive fine "
    "control. For gesture recognition, the Random Forest is backed by a simple "
    "rule-based geometric classifier. The beat stage is layered identically: a "
    "pretrained TCN beat-activation model decoded by a DBN is the default backend, "
    "and a madmom BLSTM–RNN tracker, a librosa dynamic-programming tracker "
    "[14], and finally a fixed metronome are tried in turn if a dependency or "
    "model is missing. Each layer degrades gracefully, so the game always produces "
    "a beat grid even on a machine without PyTorch or madmom."
)
heading("F.  Beat Tracking and Rhythm Generation", 2)
body(
    "The default beat tracker follows Davies and Böck [4] and is used without "
    "any training on our side. A convolutional front end feeds an eleven-layer "
    "dilated TCN whose one-dimensional beat-activation output is decoded by a DBN "
    "[4], [12] into a beat grid. We use publicly released pretrained weights [15]; "
    "our repository contains no training loop, and the checkpoint is downloaded "
    "rather than trained. A librosa [11] tempo estimate corrects occasional "
    "double-tempo errors. For each track we obtain a sequence of beat events, each "
    "a (timestamp, strength) pair, which a RhythmSpawner turns into gameplay: "
    "strong beats spawn more or stronger rocks, and beat strength scales the fall "
    "speed, so the game's tempo and intensity follow the music. When PyTorch or "
    "madmom is unavailable the system falls back to the librosa "
    "dynamic-programming tracker [11], [14], and ultimately to a fixed metronome, "
    "without changing the rest of the pipeline."
)
heading("G.  Gameplay, Scoring, and Game Design", 2)
body(
    "Three game actions are driven by the three gestures, each combined with the "
    "fingertip position. The Sword gesture slices falling rocks along the "
    "fingertip trajectory, the Fist gesture catches a moving “runner” "
    "target, and the open palm triggers a “Fever” mode that temporarily "
    "boosts scoring. Because the control signal is already stabilized upstream by "
    "the smoothing stage, these mappings respond reliably even though the raw "
    "per-frame predictions are noisy. At the end of a session the player's "
    "performance is summarized along five interpretable axes — Precision, "
    "Rhythm, Dexterity, Reaction, and Catch — and rendered as a radar chart "
    "that gives concrete per-skill feedback rather than a single opaque score."
)
heading("H.  Dynamic Difficulty Adjustment", 2)
body(
    "A dynamic difficulty-adjustment layer keeps the game in a comfortable flow "
    "state by reacting to recent behavior. It monitors the player's hits and "
    "misses over a sliding 15-second window, reduces the rock-spawn rate when the "
    "miss rate is high, and gradually restores the rate once performance "
    "stabilizes. Because the adjustment is driven by recent behavior rather than a "
    "fixed schedule, the game adapts to players of different skill levels without "
    "manual difficulty selection."
)
heading("I.  Implementation and Reproducibility", 2)
body(
    "The system targets Python 3.12 on Windows, and enabling the pretrained TCN "
    "backend required several compatibility fixes that we automate for "
    "reproducibility. Hand tracking uses the MediaPipe Tasks API with the "
    "hand_landmarker.task model. Enabling the TCN backend required pinning a CPU "
    "build of PyTorch and a Python-3.12 wheel of madmom, and patching a missing "
    "checkpoint and helper module, a deprecated librosa spectrogram call, and the "
    "removal of the collections.MutableSequence alias under Python 3.12. We vendor "
    "the checkpoint and helper code into the repository and provide a one-command "
    "setup script that installs the dependencies, applies the patches, and runs a "
    "smoke test."
)
para("Fig. 1.  System architecture. Two real-time streams — the webcam / "
     "hand-gesture pipeline and the music / rhythm pipeline — merge into a "
     "single gameplay loop with dynamic difficulty adjustment (DDA) feedback. "
     "(Insert architecture diagram here.)", italic=True, size=9)

# ================= IV. EXPERIMENTS =================
heading("IV.  Experiments and Results", 1)
heading("A.  Dataset", 2)
body(
    "We collected 807 hand samples ourselves through the webcam and stored the "
    "per-sample landmark coordinates as CSV. Although the data-collection schema "
    "reserved additional gesture labels, only three classes were ever populated: "
    "Fist with 396 samples (49.1%), Sword with 367 (45.5%), and Stop with 44 "
    "(5.5%). The resulting class imbalance is roughly 9.0× between the largest "
    "and smallest class, and the data were captured in a single environment — "
    "limitations we return to in Section V."
)
heading("B.  Gesture-Classification Accuracy", 2)
body(
    "The Random Forest is highly accurate on this three-class problem, but we read "
    "the numbers with explicit caution. On a stratified 80/20 hold-out split it "
    "reaches 99.38% accuracy, and under five-fold cross-validation it reaches "
    "98.14% ± 1.47% (Table I). The per-class F1-scores are 0.9932 (Sword), "
    "0.9937 (Fist), and 1.0000 (Stop) (Table II). Because the Stop class is "
    "classified perfectly, the residual errors are confined to the visually "
    "similar Sword and Fist classes, consistent with the PCA visualization of "
    "Section III-B. We deliberately caution against over-reading these numbers: "
    "with only three classes, fewer than 50 Stop samples, and a 9.0× "
    "imbalance, an accuracy near 99% is optimistic and is not comparable to the "
    "harder, many-class benchmarks in the gesture-recognition literature. The "
    "raw-count and row-normalized confusion matrices are shown in Fig. 2."
)
# Table I
para("Table I.  Five-fold cross-validation accuracy (mean 98.14% ± 1.47%).",
     italic=True, size=9, space_after=2)
t1 = doc.add_table(rows=1, cols=2)
t1.style = "Light Grid Accent 1"
t1.rows[0].cells[0].text = "Fold"
t1.rows[0].cells[1].text = "Accuracy (%)"
for fold, acc in [("Fold 1", "98.77"), ("Fold 2", "100.00"), ("Fold 3", "97.52"),
                  ("Fold 4", "98.76"), ("Fold 5", "95.65"), ("Mean", "98.14")]:
    row = t1.add_row().cells
    row[0].text = fold
    row[1].text = acc
doc.add_paragraph().paragraph_format.space_after = Pt(2)
# Table II
para("Table II.  Per-class F1-score (evaluation run).", italic=True, size=9, space_after=2)
t2 = doc.add_table(rows=1, cols=2)
t2.style = "Light Grid Accent 1"
t2.rows[0].cells[0].text = "Gesture"
t2.rows[0].cells[1].text = "F1-score"
for g, f in [("Sword", "0.9932"), ("Fist", "0.9937"), ("Stop", "1.0000")]:
    row = t2.add_row().cells
    row[0].text = g
    row[1].text = f
doc.add_paragraph().paragraph_format.space_after = Pt(2)

heading("C.  Real-Time Latency", 2)
body(
    "The perception pipeline is fast enough for responsive gameplay, which we "
    "verify by profiling it stage by stage (Table III). Camera capture is the "
    "dominant cost at 19.62 ms on average, while hand tracking and inference "
    "together take 13.69 ms — comparable to the 11–16 ms "
    "landmark-inference times reported for the MediaPipe Full model on mobile GPUs "
    "[1]. The end-to-end perception latency averages 33.30 ms (median 31.98 ms, "
    "p95 45.23 ms), i.e. roughly 30 frames per second. These figures were obtained "
    "in a headless benchmark with no hand in view, so the Random-Forest "
    "classification path was not exercised; the live frame rate with a hand "
    "present will differ slightly and should be reported from an on-camera run. "
    "The hand backend in this benchmark is the MediaPipe Tasks runtime using the "
    "hand_landmarker.task model."
)
para("Table III.  Per-stage latency in milliseconds (headless benchmark).",
     italic=True, size=9, space_after=2)
t3 = doc.add_table(rows=1, cols=4)
t3.style = "Light Grid Accent 1"
hdr = t3.rows[0].cells
hdr[0].text, hdr[1].text, hdr[2].text, hdr[3].text = "Stage", "Mean", "Med.", "p95"
for s, m, md, p95 in [
    ("Camera capture", "19.62", "18.67", "32.95"),
    ("Hand track + infer.", "13.69", "13.42", "17.30"),
    ("End-to-end", "33.30", "31.98", "45.23"),
]:
    row = t3.add_row().cells
    row[0].text, row[1].text, row[2].text, row[3].text = s, m, md, p95
doc.add_paragraph().paragraph_format.space_after = Pt(2)

heading("D.  Beat-Tracking Behavior", 2)
body(
    "The default TCN+DBN backend runs end-to-end and produces stable beat grids. "
    "Across four test tracks it was selected every time, with estimated tempi "
    "including 86.5 and 175.8 BPM, and the librosa tempo check corrected "
    "double-tempo cases. We verified the graceful-degradation path by removing the "
    "heavier dependencies and confirming that the librosa and metronome backends "
    "still yield a usable beat grid, so the game remains playable on machines "
    "without PyTorch or madmom."
)
para("Fig. 2.  Confusion matrices on the hold-out test set: raw counts (left) and "
     "row-normalized (right). (Insert eval_confusion_counts.png and "
     "eval_confusion_normalized.png here.)", italic=True, size=9)

# ================= V. DISCUSSION =================
heading("V.  Discussion and Limitations", 1)
body(
    "Several limitations qualify our results, and we state them explicitly. The "
    "accuracy figures reflect an easy three-class problem: the Stop class has only "
    "44 samples, the imbalance is 9.0×, and additional intended gestures were "
    "never collected, so the deployed game maps the three available classes onto "
    "its slice, catch, and Fever actions. A larger, balanced, multi-subject "
    "dataset is therefore needed before the classifier's accuracy can be "
    "considered representative, and a hold-out accuracy near 99% should be read as "
    "an upper bound on this easy task rather than a measure of general "
    "gesture-recognition difficulty. The classifier also treats each frame "
    "independently and ignores motion dynamics, so a temporal model over a window "
    "of frames is a natural next step. On the perception side, the OpenCV color "
    "fallback has low precision and is sensitive to skin tone and background, and "
    "the beat trackers — although they now run by default — are "
    "pretrained and operate offline rather than causally, so the system is not a "
    "real-time beat tracker in the strict sense."
)
body(
    "Looking further ahead, we are exploring a browser-based deployment so that "
    "the game can run without a local Python environment; candidate architectures "
    "include a WebSocket front-/back-end split and an in-browser MediaPipe-JS "
    "gesture path. We would also like to replace the rule-based and Random-Forest "
    "stages with a single trained temporal model, balance and enlarge the gesture "
    "dataset across subjects and lighting conditions, and investigate a causal TCN "
    "for genuine real-time beat tracking."
)

# ================= VI. CONCLUSION =================
heading("VI.  Conclusion", 1)
body(
    "Rockfall Riot demonstrates a fully controller-free AR rhythm game that "
    "integrates pretrained perception models with a lightweight, self-trained "
    "gesture classifier and careful systems engineering — multi-layer "
    "fallback, temporal smoothing, and dynamic difficulty adjustment. The trained "
    "component, the gesture classifier, reaches 98.1% (±1.5%) cross-validation "
    "and 99.4% hold-out accuracy on three classes — an optimistic figure that "
    "we interpret cautiously — while the full perception pipeline runs at "
    "about 30 FPS with roughly 33 ms latency. The beat tracker runs by default as "
    "a pretrained TCN+DBN model with graceful fallback to lighter algorithms. "
    "Future work includes expanding and balancing the gesture set, replacing the "
    "frame-wise classifier with a trained temporal network, adopting a causal "
    "real-time beat tracker, and deploying the game to the web."
)

# ================= REFERENCES =================
heading("References", 1)
refs = [
    "F. Zhang, V. Bazarevsky, A. Vakunov, A. Tkachenka, G. Sung, C.-L. Chang, and "
    "M. Grundmann, “MediaPipe Hands: On-device real-time hand tracking,” "
    "arXiv:2006.10214, 2020.",
    "C. Lugaresi et al., “MediaPipe: A framework for building perception "
    "pipelines,” arXiv:1906.08172, 2019.",
    "V. Bazarevsky, Y. Kartynnik, A. Vakunov, K. Raveendran, and M. Grundmann, "
    "“BlazeFace: Sub-millisecond neural face detection on mobile GPUs,” "
    "arXiv:1907.05047, 2019.",
    "M. E. P. Davies and S. Böck, “Temporal convolutional networks for "
    "musical audio beat tracking,” in Proc. 27th Eur. Signal Process. Conf. "
    "(EUSIPCO), 2019, pp. 1–5.",
    "S. Böck, F. Krebs, and G. Widmer, “A multi-model approach to beat "
    "tracking considering heterogeneous music styles,” in Proc. 15th Int. "
    "Soc. Music Inf. Retrieval Conf. (ISMIR), 2014, pp. 603–608.",
    "S. Böck, F. Krebs, and G. Widmer, “Joint beat and downbeat tracking "
    "with recurrent neural networks,” in Proc. 17th ISMIR, 2016, "
    "pp. 255–261.",
    "F. Krebs, S. Böck, and G. Widmer, “An efficient state space model "
    "for joint tempo and meter tracking,” in Proc. 16th ISMIR, 2015, "
    "pp. 72–78.",
    "S. Bai, J. Z. Kolter, and V. Koltun, “An empirical evaluation of generic "
    "convolutional and recurrent networks for sequence modeling,” "
    "arXiv:1803.01271, 2018.",
    "A. van den Oord et al., “WaveNet: A generative model for raw audio,” "
    "arXiv:1609.03499, 2016.",
    "L. Breiman, “Random forests,” Machine Learning, vol. 45, no. 1, "
    "pp. 5–32, 2001.",
    "B. McFee et al., “librosa: Audio and music signal analysis in "
    "Python,” in Proc. 14th Python in Science Conf. (SciPy), 2015, "
    "pp. 18–25.",
    "S. Böck, F. Korzeniowski, J. Schlüter, F. Krebs, and G. Widmer, "
    "“madmom: A new Python audio and music signal processing library,” "
    "in Proc. 24th ACM Int. Conf. Multimedia, 2016, pp. 1174–1178.",
    "F. Pedregosa et al., “Scikit-learn: Machine learning in Python,” "
    "J. Mach. Learn. Res., vol. 12, pp. 2825–2830, 2011.",
    "D. P. W. Ellis, “Beat tracking by dynamic programming,” J. New Music "
    "Research, vol. 36, no. 1, pp. 51–60, 2007.",
    "B. Hayes, “beat-tracking-tcn: A PyTorch implementation of the TCN beat "
    "tracker with pretrained weights,” GitHub repository, 2020. [Online]. "
    "(Verify exact title, author, and URL before submission.)",
]
for i, r in enumerate(refs, 1):
    p = doc.add_paragraph(f"[{i}] {r}")
    p.paragraph_format.space_after = Pt(2)
    p.style.font.size = Pt(9)

# ================= APPENDIX: skill deliverables =================
doc.add_page_break()
heading("Appendix A.  Claim–Evidence Map (not for submission)", 1)
body("Each major Abstract/Introduction claim is mapped to its supporting "
     "evidence, per the research-paper-writing skill's hard constraint.")
cem = [
    ("21 2.5-D landmarks via pretrained MediaPipe Hands",
     "gestures.py uses mp.solutions/tasks HandLandmarker; landmarks consumed", "supported"),
    ("807 samples, 42-D wrist-relative scale-normalized features",
     "data/gesture_data.csv = 807 rows, 43 cols; gestures.py:224 normalization", "supported"),
    ("Single-hand input",
     "app.py:180 HandTracker() -> default max_num_hands=1", "supported"),
    ("9-frame vote + 60% hysteresis + 8-frame position hold",
     "app.py:193 deque(maxlen=9); :674 ceil(len*0.60); :688 stale<=8", "supported"),
    ("Pretrained TCN+DBN default, no training loop, checkpoint downloaded",
     "rhythm.py _analyze_with_tcn imports beat_tracking_tcn; no train code in repo", "supported"),
    ("99.38% hold-out / 98.14%±1.47% CV / F1 .9932/.9937/1.0",
     "scripts/evaluate_model.py run output", "supported"),
    ("~30 FPS, ~33 ms end-to-end latency",
     "scripts/benchmark_fps.py: e2e mean 33.30 ms", "supported"),
    ("TCN selected on all 4 tracks; 86.5/175.8 BPM",
     "analyze_music(backend='tcn') run output", "supported"),
    ("Dept. = Electrical and Computer Engineering",
     "not verifiable from code", "needs author confirmation"),
]
tc = doc.add_table(rows=1, cols=3)
tc.style = "Light Grid Accent 1"
h = tc.rows[0].cells
h[0].text, h[1].text, h[2].text = "Claim", "Evidence", "Status"
for claim, ev, st in cem:
    row = tc.add_row().cells
    row[0].text, row[1].text, row[2].text = claim, ev, st

heading("Appendix B.  Five-Dimension Self-Review (not for submission)", 1)
sr = [
    ("Contribution", "Systems integration of 3 ML components into a controller-free "
     "AR loop; novelty is engineering (fallback + smoothing + DDA), not a new model. "
     "Honest about using pretrained perception/beat models. PASS with scope stated."),
    ("Writing clarity", "Each paragraph leads with its message; terminology (rock, "
     "beat event, fallback) kept stable; implementation details centralized in III-I. "
     "PASS."),
    ("Experimental strength", "Quantified gesture accuracy + latency + beat behavior. "
     "No comparison baseline for gesture task and no user study. NEEDS more evidence "
     "if targeting a research venue; acceptable for a course project."),
    ("Evaluation completeness", "Missing: ablation of smoothing parameters, "
     "per-backend beat accuracy, multi-subject data. Limitations explicitly stated. "
     "NEEDS revision for a stronger submission."),
    ("Method-design soundness", "Graceful-degradation verified; offline (non-causal) "
     "beat tracking acknowledged as a limitation; no per-case hyperparameter tuning. "
     "PASS."),
]
ts = doc.add_table(rows=1, cols=2)
ts.style = "Light Grid Accent 1"
hh = ts.rows[0].cells
hh[0].text, hh[1].text = "Dimension", "Assessment"
for dim, a in sr:
    row = ts.add_row().cells
    row[0].text, row[1].text = dim, a

doc.save(str(OUT))
print(f"Saved: {OUT}")
print(f"Paragraphs: {len(doc.paragraphs)}  Tables: {len(doc.tables)}")
